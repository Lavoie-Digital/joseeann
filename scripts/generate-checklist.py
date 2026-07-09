#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Génère un document Word (checklist) pour guider la cliente Josée-Ann Jomphe
dans la validation du contenu de son site web et la collecte des photos.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------- palette
INK      = RGBColor(0x2B, 0x28, 0x24)   # texte principal
GILT     = RGBColor(0xA8, 0x92, 0x6A)   # doré (accents)
CLAY     = RGBColor(0x8A, 0x7A, 0x63)   # sous-titres
SMOKE    = RGBColor(0x6B, 0x64, 0x5B)   # texte secondaire
LINE     = "C9BFB0"                       # bordures de tableau (hex sans #)
HEADFILL = "2B2824"                       # fond entête de tableau
ZEBRA    = "F4F0E9"                       # ligne alternée

doc = Document()

# --- styles de base ------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK

# marges
for s in doc.sections:
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.9)
    s.right_margin = Inches(0.9)


# ---------------------------------------------------------------- helpers
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color=LINE, sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def cell_text(cell, text, bold=False, color=INK, size=10, align=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def spacer(size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("")
    r.font.size = Pt(size)
    return p


def h1(text, sub=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    # petit filet doré
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = INK
    # ligne dorée
    line = doc.add_paragraph()
    line.paragraph_format.space_after = Pt(8)
    lr = line.add_run("─" * 30)
    lr.font.color.rgb = GILT
    lr.font.size = Pt(9)
    if sub:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(10)
        sr = sp.add_run(sub)
        sr.italic = True
        sr.font.size = Pt(10)
        sr.font.color.rgb = SMOKE


def eyebrow(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = GILT
    # espacement des lettres
    return p


def para(text, size=10.5, color=INK, italic=False, space_after=8, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.italic = italic
    r.bold = bold
    return p


def bullet(text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = INK
    return p


def validation_table(rows, header=("Élément à valider", "Ce qui est écrit sur le site")):
    """
    rows : liste de tuples (element, valeur_actuelle)
    Crée un tableau à 4 colonnes :
      Élément | Valeur actuelle | ☐ OK | ✍ Correction à apporter
    """
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Inches(1.9), Inches(2.7), Inches(0.6), Inches(1.9))

    hdr = table.rows[0].cells
    titles = (header[0], header[1], "✓ OK", "✍ À corriger")
    for i, (c, t) in enumerate(zip(hdr, titles)):
        set_cell_bg(c, HEADFILL)
        set_cell_borders(c)
        cell_text(c, t, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=9.5)
        c.width = widths[i]

    for idx, (el, val) in enumerate(rows):
        cells = table.add_row().cells
        cell_text(cells[0], el, bold=True, size=9.5)
        cell_text(cells[1], val, size=9.5, color=SMOKE)
        cell_text(cells[2], "☐", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(cells[3], "", size=9.5)
        for i, c in enumerate(cells):
            c.width = widths[i]
            set_cell_borders(c)
            if idx % 2 == 1:
                set_cell_bg(c, ZEBRA)
    spacer(4)
    return table


def photo_table(rows):
    """
    rows : liste de tuples (emplacement, description/suggestion)
    Colonnes : Emplacement | Photo suggérée / actuelle | ☐ Je fournis ma photo | ☐ Garder générique
    """
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = (Inches(1.9), Inches(2.7), Inches(1.35), Inches(1.15))
    hdr = table.rows[0].cells
    titles = ("Emplacement", "Ce qui est affiché actuellement",
              "☐ Je fournis\nma photo", "☐ Garder\ngénérique")
    for i, (c, t) in enumerate(zip(hdr, titles)):
        set_cell_bg(c, HEADFILL)
        set_cell_borders(c)
        cell_text(c, t, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=9)
        c.width = widths[i]
    for idx, (loc, desc) in enumerate(rows):
        cells = table.add_row().cells
        cell_text(cells[0], loc, bold=True, size=9.5)
        cell_text(cells[1], desc, size=9.5, color=SMOKE)
        cell_text(cells[2], "☐", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(cells[3], "☐", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
        for i, c in enumerate(cells):
            c.width = widths[i]
            set_cell_borders(c)
            if idx % 2 == 1:
                set_cell_bg(c, ZEBRA)
    spacer(4)
    return table


def note_box(text):
    """Encadré de note / conseil."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)
    set_cell_bg(cell, "F4F0E9")
    set_cell_borders(cell, color=GILT_HEX)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run("💡  " + text)
    r.font.size = Pt(9.5)
    r.font.color.rgb = SMOKE
    r.italic = True
    spacer(6)


GILT_HEX = "A8926A"

# ======================================================================
#                            PAGE DE TITRE
# ======================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(90)
r = title.add_run("JOSÉE-ANN JOMPHE")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = INK

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Courtier immobilier · Résidentiel & Commercial")
r.font.size = Pt(12)
r.font.color.rgb = GILT
r.italic = True

line = doc.add_paragraph()
line.alignment = WD_ALIGN_PARAGRAPH.CENTER
line.paragraph_format.space_before = Pt(20)
lr = line.add_run("─────  ✦  ─────")
lr.font.color.rgb = GILT
lr.font.size = Pt(14)

doc_title = doc.add_paragraph()
doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc_title.paragraph_format.space_before = Pt(30)
r = doc_title.add_run("Checklist de validation du site web")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = INK

doc_sub = doc.add_paragraph()
doc_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = doc_sub.add_run("Vérification des informations · Corrections · Choix des photos")
r.font.size = Pt(11)
r.font.color.rgb = SMOKE

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(120)
r = meta.add_run("Document préparé pour la révision du contenu\nÀ compléter et à retourner à votre concepteur web")
r.font.size = Pt(10)
r.font.color.rgb = CLAY

doc.add_page_break()

# ======================================================================
#                          MODE D'EMPLOI
# ======================================================================
h1("Comment utiliser ce document",
   "Quelques minutes suffisent — voici comment procéder.")

para("Ce document vous permet de passer en revue, page par page, tout le contenu "
     "de votre site web. Pour l'instant, plusieurs textes, chiffres et photos sont "
     "des exemples génériques (contenu de démonstration). Votre validation nous "
     "permettra de tout remplacer par VOS vraies informations.")

eyebrow("Trois choses à faire pour chaque section")
b = doc.add_paragraph(style="List Number"); b.add_run(
    "Vérifier chaque information : cochez la case « ✓ OK » si c'est exact.").font.size = Pt(10)
b = doc.add_paragraph(style="List Number"); b.add_run(
    "Corriger au besoin : si une information est fausse ou à changer, écrivez la "
    "bonne version dans la colonne « ✍ À corriger ».").font.size = Pt(10)
b = doc.add_paragraph(style="List Number"); b.add_run(
    "Pour les photos : indiquez si vous fournissez votre propre photo, ou si l'on "
    "garde une image générique (banque d'images professionnelle).").font.size = Pt(10)

note_box("Aucune connaissance technique requise. Écrivez directement dans le "
         "document (ou imprimez-le et annotez à la main). En cas de doute, laissez "
         "un commentaire — nous en discuterons ensemble.")

note_box("Les propriétés affichées (maisons, condos, immeubles) sont des EXEMPLES. "
         "En production, elles se synchroniseront automatiquement avec Centris® : "
         "vous n'avez donc PAS à les valider une par une ici.")

doc.add_page_break()

# ======================================================================
#            SECTION 1 — INFORMATIONS GÉNÉRALES (tout le site)
# ======================================================================
h1("1. Vos coordonnées et informations générales",
   "Ces informations apparaissent partout sur le site (bas de page, page contact, etc.). "
   "Ce sont les plus importantes à valider.")

validation_table([
    ("Nom complet affiché", "Josée-Ann Jomphe"),
    ("Titre professionnel", "Courtier immobilier agréé · Résidentiel & Commercial"),
    ("Numéro de téléphone", "514 555-0142  (numéro FICTIF à remplacer)"),
    ("Adresse courriel", "josee-ann@joseeannjomphe.ca  (à confirmer)"),
    ("Secteur desservi", "Grand Montréal & environs"),
    ("Disponibilités", "Lundi au samedi, sur rendez-vous"),
    ("Délai de réponse annoncé", "En moins de 48 heures"),
    ("Membre de l'OACIQ", "« Courtier immobilier agréé · Membre de l'OACIQ »"),
    ("Agence / bannière", "Aucune mentionnée — en ajouter une ? (ex. RE/MAX, Sutton…)"),
    ("Numéro de permis OACIQ", "Absent — à ajouter (obligatoire au Québec)"),
])

eyebrow("Réseaux sociaux")
validation_table([
    ("Page Facebook", "facebook.com/joseeannjomphecourtierimmobilier  (à confirmer)"),
    ("Compte Instagram", "instagram.com  (lien générique — à remplacer par le vôtre)"),
    ("Autre réseau (LinkedIn, TikTok…)", "Aucun — souhaitez-vous en ajouter ?"),
], header=("Réseau", "Lien actuel sur le site"))

note_box("Le numéro de téléphone et certains liens sont des exemples fictifs. "
         "Merci de fournir vos coordonnées réelles et vos liens exacts de réseaux sociaux.")

doc.add_page_break()

# ======================================================================
#                 SECTION 2 — PAGE D'ACCUEIL
# ======================================================================
h1("2. Page d'accueil",
   "La première page que vos visiteurs voient.")

eyebrow("Bannière principale (haut de page)")
validation_table([
    ("Sous-titre", "« Courtier immobilier · Résidentiel & Commercial »"),
    ("Slogan principal", "« L'immobilier d'exception mérite une signature. »"),
    ("Bouton 1", "« Découvrir les propriétés »"),
    ("Bouton 2", "« Prendre rendez-vous »"),
])
note_box("La bannière fait défiler 6 photos en arrière-plan (voir section Photos, page suivante).")

eyebrow("Section « La signature JAJ » (positionnement)")
validation_table([
    ("Texte d'accroche",
     "« Vendre ou acquérir une propriété d'exception est bien plus qu'une "
     "transaction. C'est une histoire de confiance, de timing… »"),
    ("Paragraphe descriptif",
     "« De la première visite à la signature, chaque étape est menée avec méthode… »"),
    ("Légende sous la photo", "« Josée-Ann Jomphe — Courtier immobilier agréé »"),
])

eyebrow("Section « La démarche » — les 4 étapes")
validation_table([
    ("Étape 01 — Rencontre", "« Comprendre vos objectifs, votre échéancier et vos aspirations. »"),
    ("Étape 02 — Stratégie", "« Évaluation juste, positionnement et plan de mise en marché… »"),
    ("Étape 03 — Mise en valeur", "« Photographie, staging et diffusion ciblée… »"),
    ("Étape 04 — Négociation", "« Défendre vos intérêts jusqu'à la signature, avec transparence. »"),
])

eyebrow("Section « Art de vivre & ressources »")
validation_table([
    ("Titre", "« Des guides pensés pour vous éclairer. »"),
    ("Texte", "« Vendeurs, acheteurs, amoureux des beaux lieux… »"),
])

eyebrow("Appel à l'action final")
validation_table([
    ("Titre", "« Votre prochaine adresse commence par une conversation. »"),
])

doc.add_page_break()

# ======================================================================
#                 SECTION 3 — PAGE À PROPOS
# ======================================================================
h1("3. Page « À propos »",
   "Votre présentation personnelle — probablement la section la plus personnelle à valider.")

eyebrow("Présentation")
validation_table([
    ("Titre (votre nom)", "Josée-Ann Jomphe"),
    ("Sous-titre", "Courtier immobilier agréé · Résidentiel & Commercial"),
    ("Biographie — paragraphe 1",
     "« Depuis plus de douze ans, j'accompagne des propriétaires et des "
     "acquéreurs exigeants… »   → Le nombre d'années est-il exact ?"),
    ("Biographie — paragraphe 2",
     "« Mon approche est simple : moins de bruit, plus de justesse… »"),
])

eyebrow("Vos chiffres clés (statistiques)")
validation_table([
    ("Chiffre 1", "150 M$+  « en transactions »  → exact ?"),
    ("Chiffre 2", "300+  « familles accompagnées »  → exact ?"),
    ("Chiffre 3", "12 ans  « d'expérience »  → exact ?"),
    ("Chiffre 4", "98 %  « de clients fidèles »  → exact ?"),
])
note_box("IMPORTANT : ces chiffres sont des exemples. Ils doivent refléter votre "
         "réalité. Si vous préférez ne pas afficher de chiffres, on peut retirer "
         "cette section.")

eyebrow("Vos valeurs (les 4 piliers)")
validation_table([
    ("Valeur 1 — Excellence", "« Un souci du détail qui se lit dans chaque photo, chaque visite… »"),
    ("Valeur 2 — Écoute", "« Vos objectifs guident chaque décision… »"),
    ("Valeur 3 — Intégrité", "« Transparence totale et conseils honnêtes… »"),
    ("Valeur 4 — Vision", "« Une lecture fine du marché pour saisir le bon moment… »"),
])

eyebrow("Votre citation personnelle")
validation_table([
    ("Citation",
     "« Une belle transaction, ce n'est pas seulement un prix. C'est la certitude "
     "d'avoir été bien accompagné, du premier café à la remise des clés. »"),
])
note_box("Souhaitez-vous conserver cette citation, la modifier, ou en proposer une à vous ?")

doc.add_page_break()

# ======================================================================
#                 SECTION 4 — PAGE GUIDES
# ======================================================================
h1("4. Page « Guides & art de vivre »",
   "Vos conseils pour vendeurs et acheteurs, et vos articles / inspirations.")

eyebrow("Guide du vendeur — les étapes")
validation_table([
    ("Étape 1", "Faire évaluer sa propriété à sa juste valeur marchande"),
    ("Étape 2", "Préparer et mettre en valeur les lieux (staging, photographie)"),
    ("Étape 3", "Définir une stratégie de mise en marché et de diffusion"),
    ("Étape 4", "Gérer les visites et qualifier les acheteurs sérieux"),
    ("Étape 5", "Négocier et sécuriser la meilleure offre"),
])

eyebrow("Guide de l'acheteur — les étapes")
validation_table([
    ("Étape 1", "Clarifier son budget et obtenir une préautorisation hypothécaire"),
    ("Étape 2", "Cibler les quartiers et le type de propriété recherché"),
    ("Étape 3", "Accéder aux inscriptions et aux occasions hors marché"),
    ("Étape 4", "Visiter, comparer et évaluer le potentiel réel d'un bien"),
    ("Étape 5", "Déposer une offre gagnante et bien encadrée"),
])

eyebrow("Les 3 articles / inspirations affichés")
validation_table([
    ("Article 1 — Art de vivre",
     "« Habiter un quartier : ce qui fait vraiment la valeur d'un lieu »"),
    ("Article 2 — Conseils vendeur",
     "« Cinq gestes qui augmentent la valeur perçue avant une vente »"),
    ("Article 3 — Marché",
     "« Investir dans le commercial : lire un rendement sans se tromper »"),
])
note_box("Ces 3 articles sont des exemples. Souhaitez-vous écrire vos propres articles, "
         "reprendre ces sujets, ou masquer cette section pour l'instant ?")

doc.add_page_break()

# ======================================================================
#                 SECTION 5 — PAGE CONTACT + PROPRIÉTÉS
# ======================================================================
h1("5. Page « Contact » et page « Propriétés »")

eyebrow("Page Contact")
validation_table([
    ("Titre", "« Parlons de votre projet. »"),
    ("Texte d'introduction",
     "« Une question, une évaluation, l'envie de visiter une propriété ?… »"),
    ("Coordonnées affichées",
     "Téléphone, courriel, secteur, disponibilités (voir section 1)"),
    ("Formulaire de contact",
     "Présent — vers quelle adresse courriel les messages doivent-ils arriver ?"),
])

eyebrow("Page Propriétés")
validation_table([
    ("Sous-titre", "« Une sélection rigoureuse de propriétés résidentielles et commerciales… »"),
    ("Source des propriétés",
     "Synchronisation automatique avec Centris®  → confirmer votre n° de courtier / accès"),
])
note_box("Les propriétés visibles actuellement (Westmount, Vieux-Montréal, "
         "Mont-Tremblant, etc.) sont de FAUX exemples pour illustrer le design. "
         "Une fois la synchronisation Centris® branchée, VOS vraies inscriptions "
         "s'afficheront automatiquement. Aucune action requise de votre part ici, "
         "sauf nous confirmer les accès Centris.")

doc.add_page_break()

# ======================================================================
#                 SECTION 6 — PHOTOS
# ======================================================================
h1("6. Vos photos",
   "Pour chaque emplacement, dites-nous : fournissez-vous votre propre photo, "
   "ou garde-t-on une belle image générique (banque professionnelle) ?")

eyebrow("Photo la plus importante — VOTRE portrait")
photo_table([
    ("Portrait professionnel\n(page Accueil + À propos)",
     "Une photo de vous est déjà en place (« Photo courtiere.jpg »). "
     "Est-ce la bonne ? Souhaitez-vous en fournir une nouvelle ?"),
])
note_box("Idéalement : un portrait professionnel vertical, haute résolution, "
         "sur fond neutre ou en contexte. C'est LA photo qui vous représente.")

eyebrow("Bannière d'accueil (diaporama de 6 photos)")
photo_table([
    ("Photo de couverture principale", "Grande propriété / extérieur d'exception (générique)"),
    ("Photos 2 à 6 du diaporama",
     "Mélange d'extérieurs, intérieurs et cuisines de prestige (génériques)"),
])
note_box("Recommandation : si vous avez de belles photos de propriétés que vous "
         "avez vendues (avec autorisation), elles rendront le site plus authentique. "
         "Sinon, les images génériques haut de gamme font très bien l'affaire.")

eyebrow("Autres images du site")
photo_table([
    ("Accueil — section « Art de vivre »", "Intérieur design (générique)"),
    ("Accueil — bandeau citation", "Ambiance intérieure feutrée (générique)"),
    ("Page Guides — bannière", "Intérieur haut de gamme (générique)"),
    ("Page Guides — 3 articles", "3 images d'ambiance (café, cuisine, quartier) — génériques"),
    ("Page Propriétés — bannière", "Belle propriété (générique)"),
    ("Logo / signature (JAJ)",
     "Un logo « JAJ » est utilisé. Est-ce votre logo officiel ? "
     "Fournissez la version haute résolution si vous en avez une."),
])

note_box("Comment nous envoyer vos photos : par courriel, WeTransfer ou Google Drive. "
         "Privilégiez les fichiers originaux (non compressés) pour une qualité optimale.")

doc.add_page_break()

# ======================================================================
#                 SECTION 7 — REMARQUES LIBRES
# ======================================================================
h1("7. Vos commentaires et demandes",
   "Un espace libre pour tout ce qui n'entre pas dans les cases ci-dessus.")

para("Souhaitez-vous ajouter une page, un service particulier, un témoignage de "
     "client, une section blogue, un formulaire d'évaluation en ligne ? Y a-t-il "
     "quelque chose que vous n'aimez pas dans le design actuel ? Notez tout ici :",
     space_after=10)

# lignes vides pour écrire
for _ in range(12):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("." * 95)
    r.font.color.rgb = RGBColor(0xC9, 0xBF, 0xB0)
    r.font.size = Pt(10)

spacer(14)

# --- pied de document ---
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot.paragraph_format.space_before = Pt(20)
r = foot.add_run("Merci ! Retournez ce document complété à votre concepteur web.\n"
                 "Nous mettrons le site à jour avec vos informations et vos photos.")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = CLAY

line = doc.add_paragraph()
line.alignment = WD_ALIGN_PARAGRAPH.CENTER
lr = line.add_run("─────  ✦  ─────")
lr.font.color.rgb = GILT
lr.font.size = Pt(12)

# ---------------------------------------------------------------- save
out = "/Users/zaffew/Desktop/Clients/Josee-ann/Checklist-modifications-site-web.docx"
doc.save(out)
print("Document créé :", out)
